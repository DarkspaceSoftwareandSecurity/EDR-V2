import tkinter as tk
from tkinter import scrolledtext, messagebox, ttk
import os
import threading
import time
import psutil
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler
from scapy.all import sniff
from docx import Document
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from collections import deque

# Global variables
monitoring = False
threads = []
cpu_usage = deque(maxlen=20)
memory_usage = deque(maxlen=20)
network_traffic = deque(maxlen=20)
file_changes = []
observer = None

def setup_directories():
    directories = ["PLEXOR", "EDR", "Docx-Reporting"]
    for directory in directories:
        if not os.path.exists(directory):
            os.makedirs(directory)

# Process Monitoring
def monitor_processes():
    while monitoring:
        cpu = psutil.cpu_percent()
        memory = psutil.virtual_memory().percent
        cpu_usage.append(cpu)
        memory_usage.append(memory)
        update_graphs()
        update_log(f"CPU Usage: {cpu}% | Memory Usage: {memory}%")
        time.sleep(2)

# File Monitoring
def monitor_files():
    global observer
    event_handler = FileMonitorHandler()
    observer = Observer()
    observer.schedule(event_handler, path=os.getcwd(), recursive=True)
    observer.start()
    try:
        while monitoring:
            time.sleep(1)
    except:
        observer.stop()
    observer.join()

class FileMonitorHandler(FileSystemEventHandler):
    def on_modified(self, event):
        message = f"File modified: {event.src_path}"
        file_changes.append(message)
        update_log(message)
    def on_created(self, event):
        message = f"File created: {event.src_path}"
        file_changes.append(message)
        update_log(message)
    def on_deleted(self, event):
        message = f"File deleted: {event.src_path}"
        file_changes.append(message)
        update_log(message)

# Network Monitoring
def monitor_network():
    def packet_callback(packet):
        if monitoring:
            packet_size = len(packet)
            network_traffic.append(packet_size)
            update_graphs()
            update_log(f"Captured network packet of size: {packet_size} bytes")
    sniff(prn=packet_callback, store=0, count=50, timeout=10)

# GUI Setup
root = tk.Tk()
root.title("EDR Automated Script")
root.geometry("1200x700")
root.configure(bg="#2C3E50")

# Layout Frames
top_frame = tk.Frame(root, bg="#2C3E50")
top_frame.pack(side=tk.TOP, fill=tk.X, padx=10, pady=5)

log_frame = tk.Frame(root, bg="#2C3E50")
log_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=10, pady=5)

graph_frame = tk.Frame(root, bg="#2C3E50")
graph_frame.pack(side=tk.RIGHT, fill=tk.BOTH, expand=True, padx=10, pady=5)

# Title Label
title_label = tk.Label(top_frame, text="EDR Automated Script", font=("Arial", 16, "bold"), fg="#ECF0F1", bg="#2C3E50")
title_label.pack()

# Log Area
log_area = scrolledtext.ScrolledText(log_frame, width=50, height=20, bg="#ECF0F1")
log_area.pack(fill=tk.BOTH, expand=True)

def update_log(message):
    log_area.insert(tk.END, message + "\n")
    log_area.see(tk.END)

# Matplotlib Graphs
fig, axs = plt.subplots(3, figsize=(7, 5))
canvas = FigureCanvasTkAgg(fig, master=graph_frame)
canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True)

def update_graphs():
    axs[0].cla()
    axs[1].cla()
    axs[2].cla()
    axs[0].plot(list(cpu_usage), label="CPU Usage (%)", color="blue")
    axs[1].plot(list(memory_usage), label="Memory Usage (%)", color="red")
    axs[2].plot(list(network_traffic), label="Network Traffic", color="green")
    for ax in axs:
        ax.legend()
        ax.grid(True)
    canvas.draw()

def start_edr():
    global monitoring, threads
    if monitoring:
        return
    monitoring = True
    setup_directories()
    threads = [
        threading.Thread(target=monitor_processes, daemon=True),
        threading.Thread(target=monitor_files, daemon=True),
        threading.Thread(target=monitor_network, daemon=True)
    ]
    for thread in threads:
        thread.start()
    messagebox.showinfo("EDR System", "Monitoring started...")

def stop_edr():
    global monitoring, observer
    monitoring = False
    if observer:
        observer.stop()
    messagebox.showinfo("EDR System", "Monitoring stopped...")

def generate_report():
    doc = Document()
    doc.add_heading("EDR Report", level=1)
    doc.add_heading("CPU & Memory Usage", level=2)
    doc.add_paragraph(f"CPU Usage Trend: {list(cpu_usage)}%\n")
    doc.add_paragraph(f"Memory Usage Trend: {list(memory_usage)}%\n")
    doc.add_heading("File System Activity", level=2)
    for change in file_changes:
        doc.add_paragraph(change)
    doc.add_heading("Network Traffic Analysis", level=2)
    doc.add_paragraph(f"Packet Sizes Captured: {list(network_traffic)} bytes\n")
    doc.add_paragraph("Suspicious Activity Analysis: Manual review required based on logs.")
    report_path = os.path.join(os.getcwd(), "EDR_Report.docx")
    doc.save(report_path)
    messagebox.showinfo("Report Generated", f"Report saved at: {report_path}")

button_frame = tk.Frame(root, bg="#2C3E50")
button_frame.pack(side=tk.BOTTOM, fill=tk.X, padx=10, pady=5)

tk.Button(button_frame, text="Start Monitoring", command=start_edr, bg="#3498DB", fg="white").pack(side=tk.LEFT, padx=5, pady=5)
tk.Button(button_frame, text="Stop Monitoring", command=stop_edr, bg="#E74C3C", fg="white").pack(side=tk.LEFT, padx=5, pady=5)
tk.Button(button_frame, text="Generate Report", command=generate_report, bg="#2ECC71", fg="white").pack(side=tk.LEFT, padx=5, pady=5)
tk.Button(button_frame, text="Exit", command=root.quit, bg="#95A5A6", fg="white").pack(side=tk.RIGHT, padx=5, pady=5)

root.mainloop()
